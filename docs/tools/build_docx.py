"""Builds docs/chapters/*.md into docs/FYP_Report.docx (the full report)
and, optionally, a standalone DOCX per chapter (docs/FYP_Report_ChapterN_*.docx)
for sharing a single chapter with a supervisor ahead of the full report
being finished. Uses pandoc (pypandoc_binary), substituting the 4 ASCII
sketches for the diagram PNGs (run gen_figure_pngs.py first), then
post-processes with python-docx to: force each chapter and preliminary page
onto its own page, keep every heading glued to the content right after it,
keep short sub-sections from splitting across a page boundary, keep tables
from splitting across a page boundary, center the Title Page, and (full
report only) split the document into two independently-numbered sections
matching the department's format: the preliminary pages (Title through
Abstract) numbered i, ii, iii... in lower-case roman numerals, and the main
work (Chapter One onward) numbered 1, 2, 3... starting fresh at 1.

A standalone chapter DOCX has no preliminary pages and is numbered 1, 2,
3... on its own (Word's default), independent of where that chapter will
eventually land in the full report.

Requires: pip install pypandoc_binary python-docx (into a venv is fine).

Usage:
    python3 gen_figure_pngs.py
    python3 build_docx.py                 # full report only (default)
    python3 build_docx.py --chapters      # full report + one DOCX per chapter
    python3 build_docx.py --chapters-only # one DOCX per chapter, skip the full report
"""
import argparse
import copy
import re
from pathlib import Path

import pypandoc
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

HERE = Path(__file__).parent
CHAPTERS_DIR = HERE.parent / 'chapters'
DOCX_PATH = HERE.parent / 'FYP_Report.docx'

CHAPTER_ONE_MARKER = 'CHAPTER ONE'
KEEP_BLOCK_MAX_CHARS = 4000

FENCE_RE = re.compile(r'```\n.*?\n```', re.S)
PNGS = [HERE / 'fig_2_1.png', HERE / 'fig_3_1.png', HERE / 'fig_3_2.png', HERE / 'fig_4_1.png']
for p in PNGS:
    assert p.exists(), f'{p} missing - run gen_figure_pngs.py first'


def chapter_paths():
    return sorted(p for p in CHAPTERS_DIR.glob('*.md') if not p.name.startswith('00_'))


def prelim_path():
    matches = list(CHAPTERS_DIR.glob('00_*.md'))
    assert len(matches) == 1, f'expected exactly one 00_*.md preliminary-pages file, found {len(matches)}'
    return matches[0]


def figure_offset_for(path):
    offset = 0
    for p in chapter_paths():
        if p == path:
            return offset
        offset += len(FENCE_RE.findall(p.read_text()))
    raise ValueError(f'{path} not found among chapter files')


def substitute_figures(src, pngs_slice):
    idx = 0

    def replace_block(m):
        nonlocal idx
        png = pngs_slice[idx]
        idx += 1
        return f'![]({png})'

    new_src = FENCE_RE.sub(replace_block, src)
    assert idx == len(pngs_slice), f'expected {len(pngs_slice)} code blocks, replaced {idx}'
    return new_src


def add_grid_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '888888')
        borders.append(el)
    tblPr.append(borders)


def prevent_row_split(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)


def insert_section_break_before(doc, paragraph):
    """Ends the current (single) section at a new empty paragraph inserted
    just before `paragraph`, by giving that new paragraph a copy of the
    document's sectPr. The document's original trailing sectPr becomes
    section 2's properties, unchanged."""
    body = doc.element.body
    orig_sectPr = body.find(qn('w:sectPr'))
    new_p = OxmlElement('w:p')
    paragraph._p.addprevious(new_p)
    pPr = OxmlElement('w:pPr')
    new_p.append(pPr)
    pPr.append(copy.deepcopy(orig_sectPr))


def set_page_number_format(section, fmt, start=1):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    pgNumType.set(qn('w:start'), str(start))


def add_page_number_footer(section):
    section.footer.is_linked_to_previous = False
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for run in list(p.runs):
        run._r.getparent().remove(run._r)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def build(src, out_path, with_prelim):
    temp_md = HERE / f'_{out_path.stem}_for_docx.md'
    temp_md.write_text(src)
    pypandoc.convert_file(str(temp_md), 'docx', outputfile=str(out_path), extra_args=['--standalone'])
    temp_md.unlink()

    doc = Document(str(out_path))

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)

    for level, size in [(1, 16), (2, 13), (3, 12)]:
        try:
            style = doc.styles[f'Heading {level}']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
        except KeyError:
            pass

    doc.styles['Heading 1'].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    paras = doc.paragraphs
    prev_style = None
    first_heading_seen = False
    title_page_para = None
    chapter_one_para = None

    for p in paras:
        style_name = p.style.name if p.style else ''
        is_heading = style_name.startswith('Heading')
        if is_heading:
            p.paragraph_format.keep_with_next = True
            if style_name == 'Heading 1':
                text = p.text.strip()
                if text == 'TITLE PAGE':
                    title_page_para = p
                if text == CHAPTER_ONE_MARKER:
                    chapter_one_para = p
                if not first_heading_seen:
                    first_heading_seen = True
                elif prev_style == 'Heading 1':
                    pass
                else:
                    p.paragraph_format.page_break_before = True
        if p.runs and p.runs[0].bold and p.text.strip().startswith('Figure '):
            p.paragraph_format.keep_with_next = True
        prev_style = style_name

    # center the Title Page block (everything from "TITLE PAGE" up to the next heading)
    if title_page_para is not None:
        started = False
        for p in paras:
            if p is title_page_para:
                started = True
                continue
            if started:
                style_name = p.style.name if p.style else ''
                if style_name.startswith('Heading'):
                    break
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # keep each short (< KEEP_BLOCK_MAX_CHARS) h2/h3 sub-section from splitting
    # across a page boundary, by chaining keep_with_next through its paragraphs;
    # a section too long to plausibly fit one page, or one containing a table,
    # is left to flow and split naturally rather than getting pushed whole onto
    # a fresh page and wasting the one it left mostly blank.
    body_children = list(doc.element.body)
    table_positions = {id(t._tbl): body_children.index(t._tbl) for t in doc.tables}

    groups = []
    current_start = None
    for i, p in enumerate(paras):
        style_name = p.style.name if p.style else ''
        if style_name in ('Heading 2', 'Heading 3'):
            if current_start is not None:
                groups.append((current_start, i))
            current_start = i
        elif style_name == 'Heading 1':
            if current_start is not None:
                groups.append((current_start, i))
            current_start = None
    if current_start is not None:
        groups.append((current_start, len(paras)))

    for start, end in groups:
        start_pos = body_children.index(paras[start]._p)
        end_pos = body_children.index(paras[end - 1]._p) if end - 1 < len(paras) else len(body_children)
        contains_table = any(start_pos <= tp <= end_pos for tp in table_positions.values())
        text_len = sum(len(paras[j].text) for j in range(start, end))
        if not contains_table and text_len <= KEEP_BLOCK_MAX_CHARS:
            for j in range(start, end - 1):
                paras[j].paragraph_format.keep_with_next = True

    for table in doc.tables:
        add_grid_borders(table)
        prevent_row_split(table)
        n_rows = len(table.rows)
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.keep_together = True
                    if i < n_rows - 1:
                        p.paragraph_format.keep_with_next = True

    if with_prelim:
        # split into two independently-numbered sections at Chapter One
        assert chapter_one_para is not None, 'could not find the "CHAPTER ONE" heading'
        chapter_one_para.paragraph_format.page_break_before = False
        insert_section_break_before(doc, chapter_one_para)

        prelim_section, body_section = doc.sections[0], doc.sections[1]
        set_page_number_format(prelim_section, 'lowerRoman', start=1)
        set_page_number_format(body_section, 'decimal', start=1)
        add_page_number_footer(prelim_section)
        add_page_number_footer(body_section)
    else:
        add_page_number_footer(doc.sections[0])

    doc.save(str(out_path))
    print('wrote', out_path)


def build_full():
    src = '\n\n---\n\n'.join(p.read_text() for p in [prelim_path(), *chapter_paths()])
    src = substitute_figures(src, PNGS)
    build(src, DOCX_PATH, with_prelim=True)


def build_chapter(path):
    src = path.read_text()
    offset = figure_offset_for(path)
    n = len(FENCE_RE.findall(src))
    src = substitute_figures(src, PNGS[offset:offset + n])
    label = path.stem.split('_', 1)[1]  # e.g. 'chapter_one_introduction'
    out_path = HERE.parent / f'FYP_Report_{label}.docx'
    build(src, out_path, with_prelim=False)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--chapters', action='store_true', help='also build a standalone DOCX per chapter')
    parser.add_argument('--chapters-only', action='store_true', help='build only the standalone chapter DOCX files, skip the full report')
    args = parser.parse_args()

    if not args.chapters_only:
        build_full()
    if args.chapters or args.chapters_only:
        for p in chapter_paths():
            build_chapter(p)


if __name__ == '__main__':
    main()
